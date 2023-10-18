#BY Gerard Mulvany
import os
import tempfile
from tkinter.font import BOLD
import pypandoc 
import re
from docx import Document
from lxml import etree
from xml.etree.ElementTree import QName


####################### GLOBAL PROPERTIES ############################
replace_video_links = False
use_apa_figure_style = False

current_directory = os.getcwd()
																		 
doc_path = current_directory

# ANSI escape codes for console red text
bold_red = "\033[1;31m"
green = "\033[32m"
yellow = "\033[93m"
italic = "\033[3m"
reset = "\033[0m"

#Pandoc arguments
extra_args = [
			   '--filter',
			   'pandoc-fignos',
			   '--citeproc',
			   '--csl',
			   'custom-key.csl',
			   '--biblatex',
			   '--bibliography',
			   'Bibliography.bib',
			   '--wrap',
			   'none',
			   '--extract-media',
			   './','--verbose']
to_format = 'tex'
input_format = 'docx+citations'

####################################### MAIN FUNCTIONS ################################################################
		  
def generate_tex(input_file):
	try:
		#update docx for current chapter number
		doc_file = input_file
		tex_file = input_file + ".tex"
	except Exception as e:
		print(bold_red,"Error: Find document failed :", e , reset)
		return False
	  
	if not os.path.exists(doc_file):
		print(bold_red + "Error: Missing file path: " + doc_file + reset)
		return False
	else:
		print("Processing File:", doc_file)
		try:
				conversion_result = convert_docx_to_tex(doc_file, tex_file)
				if conversion_result:
					try:
							fixes_result = manual_fixes(input_file)
							return fixes_result
					except:
							print(bold_red + "Error: Manual Fixes failed" + reset)
		except :
				print(bold_red + "Error: Attempted but failed to convert document" + reset)
	return False


def convert_docx_to_tex(input_file, output_file):
	try:
		pypandoc.convert_file(input_file, to_format, input_format, outputfile=output_file, extra_args=extra_args)
		return True
	except Exception as e:
		print(bold_red, "Pandoc ERROR:", str(e), reset)
		return False

def manual_fixes(input_file):
	#create and verify zip version of document then unzip to temp
	doc_file = input_file
	tex_file = input_file + ".tex"
	result = False
	captions = []
	image_names = []
	subdoc_captions = []
	had_error = False
	pattern = r'^\\includegraphics\[width=.*in,height=.*in\]{'
	try:
		image_names = get_original_image_names(doc_file)    
	except Exception as e:
		print(bold_red, "Failed to find image names:", e, reset)
		had_error = True

	try: #Copy each caption for figures into their own temporary docx so they can be individually pandoc converted and cross-references can be linked
		subdoc_captions = extract_paragraphs_by_style(doc_file, "Caption")
	except Exception as e:
		print(bold_red, "Failed to find captions by style:", e, reset)
		had_error = True

	try:
		crossrefs = find_text_with_style(doc_file, "CrossReference")
	except Exception as e:
		print(bold_red, "Failed to find cross-references by style:", e, reset)
		had_error = True

	try: #remove empty captions          
		for c in captions:
			if c == "":
				print(yellow, "Warning: Removed empty caption", reset)
				had_error = True
				captions.remove(c)
	except Exception as e:
		print(bold_red, "Remove empty captions failed:", e, reset)
		had_error = True

	#Console line to inform
	print(green, "Number of Images = ", len(image_names))
	print(green, "Number of Captions = ", len(subdoc_captions), reset)
	if len(image_names) != len(subdoc_captions):
		print(bold_red, "ERROR: Figures and captions mismatch", reset)
		had_error = True
	for name in image_names:
		i = image_names.index(name)
		print(yellow, f"Progress: {i}/{len(image_names)}", reset, end='\r')
		if use_apa_figure_style:
			make_apa7_figure(name, i, tex_file, subdoc_captions)
		
		
		delimiter = "Note."
		before_note = ""
		#search_pattern = before_note.encode('utf-8', 'escape')
		after_note = ""
		end_search = ""
		saved_note = ""
		has_note = False
		latex_caption = ""
		cap_search = ""
		search_pattern = b""

		try : #Remove old end figure command
			replace_line_with_pattern(tex_file, r"\\end{fignos:no-prefix-figure-caption}", "")
			result = True
		except:
			print(bold_red, "Error changing to \\end{figure}", reset)
			had_error = True

		try: #Get OS path for image
			name = os.path.basename(name)
		except Exception as e:
			print(bold_red, "ERROR: Failed to get OS path for Image: ", name, "   Error is:", e, reset)
			had_error = True

		try: #replace figure 'begin' environment
			replace_line_with_pattern(tex_file, r"\\begin{fignos:no-prefix-figure-caption}", r"\begin{figure}")
			result = True
		except Exception as e:
			print(bold_red, "Error changing \\begin{figure} WITH ERROR:", e, reset)
			had_error = True

		label_name, *_ = name.split(".")
		label_name = r"    \label{fig:" + label_name + "}"
		replacement = "    \\includegraphics[width=\\textwidth]{" + "Figures/" + name + "}"
		try:#Make figures text width and add image filepath
			replace_line_with_pattern(tex_file, pattern, replacement)
			result = True    
		except Exception as e:
			print(bold_red, "Error adjusting \\includegraphics: ", e, reset)
			had_error = True

		if replace_video_links:
			create_video_icons(tex_file)
		try:          

			latex_caption = pypandoc.convert_file(subdoc_captions[i], to_format, input_format, extra_args=extra_args)
			latex_caption = replace_caption_cross_references(latex_caption)
			if delimiter in latex_caption:
				before_note, after_note = latex_caption.split(delimiter, 1)
				before_note = r"    \caption{" + before_note + "}"

				after_note = r"\textit{" + delimiter + " }" + after_note.strip() # Include the delimiter in the "after_note" string
				after_note = r"" + after_note
				
				after_note = "    " + r"    \raggedright{\small{" + after_note + "}}"
				#search_pattern = get_first_4_words(latex_caption)
				search_pattern = truncate_and_encode(latex_caption, 50)
				saved_note = replace_line_with_pattern(tex_file, search_pattern, after_note, True) #Replace the original under-image caption with only the "after-note" section
				has_note = True
			else:
				before_note = r"" + latex_caption
				before_note = before_note.strip()
				before_note = r"    \caption{" + before_note + "}"
				caption_replacement = ""
				#search_pattern = get_first_4_words(latex_caption)
				search_pattern = truncate_and_encode(latex_caption, 50)
				replace_line_with_pattern(tex_file, search_pattern, caption_replacement) #Delete original caption that was under the image
				has_note = False

			add_new_line_of_text_above_word(tex_file, replacement, before_note) #Add the figure title/short caption above image
			result = True
		except Exception as e:
			print(bold_red, "Failed to correct caption:", e, reset)
			had_error = True
		try:      
			cap_search =r"\\caption{" + search_pattern.decode("utf-8", 'ignore')  #Ignore unbound 
			#time.sleep(0.1)
			#add_line_below_pattern(tex_file, cap_search, label_name) #Add a label for cross-referencing
			add_new_line_of_text_above_word(tex_file, replacement, label_name)
			
			result = True
		except Exception as e:
			print(bold_red, "Error adding label:", e, reset)
			had_error = True
		try:
			if cap_search != "":
				#add_line_below_pattern(tex_file, cap_search, r"    \centering") #Centre the image
				add_new_line_of_text_above_word(tex_file, label_name, r"    \centering")
		except Exception as e:
			print(bold_red, "Error adding centering cmd:", e, reset)
		try: 
			if has_note:
				end_search =  saved_note
			else:
				end_search = replacement 
			add_new_line_of_text_below_word(tex_file, end_search, r"\end{figure}") #Close the figure environment
			result = True
		except Exception as e:
			print(bold_red, "Error adding \\end{Figure} line:", e, reset)
			had_error = True
	
	        
	try:
		replace_cross_references(tex_file) #Run the function that replaces cross refs based on style with LATEX version
	except Exception as e:
		print(bold_red, "Replace captions failed:", e, reset)
		had_error = True

	if len(image_names) == 0 and had_error == False:
		return True
	return result

def make_apa7_figure(name, i, tex_file, subdoc_captions):
	delimiter = "Note."
	before_note = ""
	#search_pattern = before_note.encode('utf-8', 'escape')
	after_note = ""
	end_search = ""
	saved_note = ""
	has_note = False
	latex_caption = ""
	cap_search = ""
	search_pattern = b""
	pattern = r'^\\includegraphics\[width=.*in,height=.*in\]{'

	try : #Remove old end figure command
		replace_line_with_pattern(tex_file, r"\\end{fignos:no-prefix-figure-caption}", "")
		result = True
	except:
		print(bold_red, "Error changing to \\end{figure}", reset)
		had_error = True

	try: #Get OS path for image
		name = os.path.basename(name)
	except Exception as e:
		print(bold_red, "ERROR: Failed to get OS path for Image: ", name, "   Error is:", e, reset)
		had_error = True

	try: #replace figure 'begin' environment
		replace_line_with_pattern(tex_file, r"\\begin{fignos:no-prefix-figure-caption}", r"\begin{figure}")
		result = True
	except Exception as e:
		print(bold_red, "Error changing \\begin{figure} WITH ERROR:", e, reset)
		had_error = True

	label_name, *_ = name.split(".")
	label_name = r"    \label{fig:" + label_name + "}"
	replacement = "    \\includegraphics[width=\\textwidth]{" + "Figures/" + name + "}"
	try:#Make figures text width and add image filepath
		replace_line_with_pattern(tex_file, pattern, replacement)
		result = True    
	except Exception as e:
		print(bold_red, "Error adjusting \\includegraphics: ", e, reset)
		had_error = True
		
	if replace_video_links:
		create_video_icons(tex_file)
	try:          

		latex_caption = pypandoc.convert_file(subdoc_captions[i], to_format, input_format, extra_args=extra_args)
		latex_caption = replace_caption_cross_references(latex_caption)
		if delimiter in latex_caption:
			before_note, after_note = latex_caption.split(delimiter, 1)
			before_note = r"    \caption{" + before_note + "}"

			after_note = r"\textit{" + delimiter + " }" + after_note.strip() # Include the delimiter in the "after_note" string
			after_note = r"" + after_note
			
			after_note = "    " + r"    \raggedright{\small{" + after_note + "}}"
			#search_pattern = get_first_4_words(latex_caption)
			search_pattern = truncate_and_encode(latex_caption, 50)
			saved_note = replace_line_with_pattern(tex_file, search_pattern, after_note, True) #Replace the original under-image caption with only the "after-note" section
			has_note = True
		else:
			before_note = r"" + latex_caption
			before_note = before_note.strip()
			before_note = r"    \caption{" + before_note + "}"
			caption_replacement = ""
			#search_pattern = get_first_4_words(latex_caption)
			search_pattern = truncate_and_encode(latex_caption, 50)
			replace_line_with_pattern(tex_file, search_pattern, caption_replacement) #Delete original caption that was under the image
			has_note = False

		add_new_line_of_text_above_word(tex_file, replacement, before_note) #Add the figure title/short caption above image
		result = True
	except Exception as e:
		print(bold_red, "Failed to correct caption:", e, reset)
		had_error = True
	try:      
		cap_search =r"\\caption{" + search_pattern.decode("utf-8", 'ignore')  #Ignore unbound 
		#time.sleep(0.1)
		#add_line_below_pattern(tex_file, cap_search, label_name) #Add a label for cross-referencing
		add_new_line_of_text_above_word(tex_file, replacement, label_name)
		
		result = True
	except Exception as e:
		print(bold_red, "Error adding label:", e, reset)
		had_error = True
	try:
		if cap_search != "":
			#add_line_below_pattern(tex_file, cap_search, r"    \centering") #Centre the image
			add_new_line_of_text_above_word(tex_file, label_name, r"    \centering")
	except Exception as e:
		print(bold_red, "Error adding centering cmd:", e, reset)
	try: 
		if has_note:
			end_search =  saved_note
		else:
			end_search = replacement 
		add_new_line_of_text_below_word(tex_file, end_search, r"\end{figure}") #Close the figure environment
		result = True
	except Exception as e:
		print(bold_red, "Error adding \\end{Figure} line:", e, reset)
		had_error = True

def make_acm_figure(name, i, tex_file, subdoc_captions):
	saved_note = ""
	pattern = r'^\\includegraphics\[width=.*in,height=.*in\]{'
	try :
		replace_line_with_pattern(tex_file, r"\\end{fignos:no-prefix-figure-caption}", "")
		had_error = True
	except Exception as e:
		print(bold_red, e, reset)
	label_name, *_ = name.split(".")
	label_name = r"    \label{fig:" + label_name + "}"
	replacement = "    \\includegraphics[width=\\textwidth]{" + "Figures/" + name + "}"
	try:#Make figures text width and add image filepath
		replace_line_with_pattern(tex_file, pattern, replacement)
		had_error = True    
	except Exception as e:
		print(bold_red, "Error adjusting \\includegraphics: ", e, reset)
		had_error = True
		
	latex_caption = pypandoc.convert_file(subdoc_captions[i], to_format, input_format, extra_args=extra_args)
	latex_caption = replace_caption_cross_references(latex_caption)
	latex_caption = r"    \caption{" + latex_caption + "}"
	add_new_line_of_text_above_word(tex_file, latex_caption, label_name)
	search_pattern = truncate_and_encode(latex_caption, 50)
	saved_note = replace_line_with_pattern(tex_file, search_pattern, latex_caption, True) #Replace the original under-image caption with only the "after-note" section
	add_new_line_of_text_below_word(tex_file, saved_note, r"\end{figure}") #Close the figure environment


def create_video_icons(tex_file):
	try:
		modified_text, unedited_line = replace_youtube_link_with_command(tex_file, 'Youtube')
		if modified_text is not None:
			add_new_line_of_text_above_word(tex_file, unedited_line.decode('utf-8', 'replace'), modified_text.decode('utf-8', 'replace')[:-1])
	except Exception as e:
		print(bold_red, e, reset)
	try:
		modified_text, unedited_line = replace_youtube_link_with_command(tex_file, 'Vimeo')
		if modified_text is not None:
			add_new_line_of_text_above_word(tex_file, unedited_line.decode('utf-8', 'replace'), modified_text.decode('utf-8', 'replace')[:-1])
	except Exception as e:
		print(bold_red, e, reset)

###################################### SHORT FUNCTIONS ###################################################################

def sync_chapter_to_overleaf(project_ref, chapter_num):                  
	new_tex = update_file_number(chapter_num, ".tex")
	old_tex = "Chapter" + chapter_num + ".tex"
	project_ref.update_file(old_tex, new_tex)  


def update_file_number(chapter_number, extension):
	file = os.path.join(doc_path, 'Chapter')
	file = file + str(chapter_number) + extension
	return file

def get_original_image_names(docx_file):
	doc = Document(docx_file)
	image_names = []
	refs = []
	ordered_images = []
	all_rids = []

	try:
		for rel in doc.part.rels.values():
			if "image" in rel.reltype:
				try:
					image_name = rel._target.split("/")[-1]
					image_name = r"" + image_name
					image_names.append(image_name)
					ref = rel.rId 
					#print(ref)
					refs.append(ref)
				except:
					pass
	except Exception as e:
		print(bold_red, "rel.values failed:", e, reset)
	try:
		all_rids = get_rid_order(docx_file)
	except Exception as e:
		print(bold_red, "All rID's failed:",e, reset)
	try:
		for rid in all_rids:
			if rid in refs:
				pass
			else:
				all_rids.remove(rid)
	except Exception as e:
		print(bold_red, "Fix all rID list failed:",e, reset)

	#reordered_rids = sorted(refs, key=lambda x: all_rids.index(x))
	try:
		ordered_images = [c for _, c in sorted(zip(refs, image_names), key=lambda x: all_rids.index(x[0]))]
	except Exception as e:
		print(bold_red, "order_images Failed:", e, reset)
	return ordered_images

def remove_trailing_whitespace(string):
	return string.rstrip()

def get_first_4_words(string):
	words = string.split()
	first_4_words = words[:4]
	return ' '.join(first_4_words)

def get_last_4_words(string):
	words = string.split()
	last_4_words = words[-4:]
	return ' '.join(last_4_words)

def truncate_and_encode(input_string, max_length=30):
	# Check if the input string length is less than or equal to max_length
	if len(input_string) <= max_length:
		# Encode the entire string to UTF-8 with escape characters
		return input_string.encode('utf-8', 'escape')
	else:
		# Truncate the string to the first 30 characters and encode to UTF-8 with escape characters
		return input_string[:max_length].encode('utf-8', 'escape')

def add_line_above_first_line(tex_file, new_line):
	with open(tex_file, 'r', encoding='utf-8') as file:
		lines = file.readlines()

	lines.insert(0, new_line + '\n')

	with open(tex_file, 'w', encoding='utf-8') as file:
		file.writelines(lines)

def find_replace_unknown(original, replacement):
	# Escape special characters in the original string
	escaped_original = re.escape(original)
	
	# Create a regular expression pattern with the escaped original string
	pattern = re.compile(escaped_original)
	
	# Replace the unknown characters in the replacement string using the pattern
	modified = re.sub(pattern, replacement, original)
	
	return modified

def replace_cross_references(file_path):
	pattern = r'\\emph{\\textbf{\\ul{([^{}]+)}}}'

	with open(file_path, 'r', encoding='utf-8') as file:
		content = file.read()

	def repl(match):
		cross_ref = match.group(1)
		cross_ref = cross_ref.replace("\\", "")
		return r'\ref{' + cross_ref + '}'

	modified_content = re.sub(pattern, repl, content)

	with open(file_path, 'w', encoding='utf-8') as file:
		file.write(modified_content)

def replace_first_occurrence(file_path, word_to_replace, replacement):
	with open(file_path, 'r', encoding='utf-8') as file:
		content = file.read()

	# Find the index of the first occurrence of the word
	index = content.find(word_to_replace)

	if index != -1:
		# Replace the word with the desired replacement
		updated_content = content[:index] + replacement + content[index + len(word_to_replace):]

		# Write the updated content back to the file
		with open(file_path, 'w', encoding='utf-8') as file:
			file.write(updated_content)

def replace_line_with_pattern(file_path, pattern, replacement, is_after_caption=False):
	replaced = False  # Flag to track if replacement has been made
	delimiter = 'Note.'
	before_note = ''
	after_note = ''

	# Check if 'pattern' is of type bytes
	if isinstance(pattern, bytes):
		pattern = pattern.decode('utf-8')

	with open(file_path, 'r', encoding='utf-8') as file:
		lines = file.readlines()

	with open(file_path, 'w', encoding='utf-8') as file:
		for line in lines:
			if not replaced:
				if pattern in line or re.search(pattern, line):
					if is_after_caption:
						before_note, after_note = line.split(delimiter, 1)
						after_note = r"\textit{" + delimiter + " }" + after_note.strip()  # Include the delimiter in the "after_note" string
						after_note = r"" + after_note
						after_note = r"    \raggedright{\small{" + replace_caption_cross_references(after_note) + "}}"
						file.write(after_note + '\n')
					else:
						file.write(replacement + '\n')
					replaced = True  # Set the flag to True after making the replacement
				else:
					file.write(line)
			else:
				file.write(line)
	return after_note

def add_line_above_pattern(file_path, pattern, new_line):
	with open(file_path, 'r', encoding='utf-8') as file:
		content = file.read()

	modified_content = re.sub(pattern, new_line + r'\n\g<0>', content)

	with open(file_path, 'w', encoding='utf-8') as file:
		file.write(modified_content)

def add_line_below_pattern(file_path, pattern, new_line):
	with open(file_path, 'rb') as file:
		content = file.read()

	# Ensure that the pattern is a bytes object
	if isinstance(pattern, str):
		pattern = pattern.encode('utf-8')

	# Use bytes pattern to search in content
	pattern = re.compile(pattern)
	modified_content = pattern.sub(lambda x: x.group(0) + b'\n' + new_line.encode('utf-8'), content, count=1)

	with open(file_path, 'wb') as file:
		file.write(modified_content)



def find_text_with_style(document_path, style_name):
	doc = Document(document_path)
	text_with_style = []

	for paragraph in doc.paragraphs:
		if paragraph.style.name == style_name:
			text_with_style.append(paragraph.text)

	return text_with_style

def find_and_replace(file_path, search_text, replace_text):
	with open(file_path, 'r', encoding='utf-8') as file:
		content = file.read()

	try:
		modified_content = content.replace(search_text, replace_text)
		with open(file_path, 'w', encoding='utf-8') as file:
			file.write(modified_content)
	except Exception as e:
		print(bold_red, "Find and Replace Failed:", e, reset)

def get_rid_order(docx_file):
	# Open the document using python-docx
	doc = Document(docx_file)

	# Retrieve the document.xml part
	document_part = doc.part.document

	# Convert the document XML blob to string
   # xml_string = document_part.part.blob.decode('utf-8')
	root = None
	drawing_elements = []
	namespaces = {}
	try:
		# Parse the document XML using lxml
		root = etree.fromstring(doc.part.blob, parser=None)
	except Exception as e:
		print(bold_red, "XML Parsing Failed:", e, reset)

	try:
	# Register the 'w' namespace prefix
		namespaces = {
			'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
			'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
			'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
			'cx1': 'http://schemas.microsoft.com/office/drawing/2015/9/8/chartex',
			'cx2': 'http://schemas.microsoft.com/office/drawing/2015/10/21/chartex',
			'cx3': 'http://schemas.microsoft.com/office/drawing/2016/5/9/chartex',
			'cx4': 'http://schemas.microsoft.com/office/drawing/2016/5/10/chartex',
			'cx5': 'http://schemas.microsoft.com/office/drawing/2016/5/11/chartex',
			'cx6': 'http://schemas.microsoft.com/office/drawing/2016/5/12/chartex',
			'cx7': 'http://schemas.microsoft.com/office/drawing/2016/5/13/chartex',
			'cx8': 'http://schemas.microsoft.com/office/drawing/2016/5/14/chartex',
			'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
			'aink': 'http://schemas.microsoft.com/office/drawing/2016/ink',
			'am3d': 'http://schemas.micro...17/model3d',
			'o': 'urn:schemas-microsoft-com:office:office',
			'oel': 'http://schemas.microsoft.com/office/2019/extlst',
			'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
			'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
			'v': 'urn:schemas-microsoft-com:vml',
			'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
			'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
			'w10': 'urn:schemas-microsoft-com:office:word',
			'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
			'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
			'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
			'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
			'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
			'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
			'w16du': 'http://schemas.microsoft.com/office/word/2023/wordml/word16du',
			'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
			'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
			'wpg': 'http://schemas.microsoft.com/office/' }

			
		# Find all w:drawing elements in the document.xml
		#drawing_elements = root.findall(".//w:drawing", namespaces=namespaces)
		if root != None:      
			drawing_elements = root.findall(".//a:blip", namespaces=namespaces)
		else:
			print(bold_red, "Get RID Order Failed: doc.part.blob.root is None", reset)

	except Exception as e:
		print(bold_red, "w:drawing search failed:", e, reset)

	# Retrieve the rId values in the order they appear in the document.xml
	rid_order = []
	try:
		for drawing_element in drawing_elements:
			#blip_element = drawing_element.find(".//a:blip", namespaces=namespaces)
			if drawing_element is not None:
				embed_attrib_name = QName(namespaces['r'], 'link')
				try:
					str_attrib = str(embed_attrib_name)
					r_id = drawing_element.attrib[str_attrib]
					rid_order.append(r_id)
				except KeyError:
					print(bold_red,"Failed to get image link:", KeyError, reset)
	except Exception as e:
		print(bold_red, "rid list failed with a non-key-error:", e)
	return rid_order


def add_new_line_of_text_above_word(tex_file, specific_word, new_line_text):
	# Read the content of the tex file
	with open(tex_file, 'r', encoding='utf-8') as file:
		content = file.readlines()

	# Iterate over the lines and add a new line of text above the first line with the specific word
	modified_content = []
	word_encountered = False  # Flag variable to track if the specific word has been encountered
	for line in content:
		if specific_word in line and not word_encountered:
			modified_content.append(new_line_text + '\n')  # Add the new line of text
			word_encountered = True  # Set the flag to True after encountering the word
		modified_content.append(line)

	# Write the modified content back to the tex file
	with open(tex_file, 'w', encoding='utf-8') as file:
		file.writelines(modified_content)

def add_new_line_of_text_below_word(tex_file, specific_word, new_line_text):
	# Read the content of the tex file
	with open(tex_file, 'r', encoding='utf-8') as file:
		content = file.readlines()

	# Iterate over the lines and add a new line of text above the first line with the specific word
	modified_content = []
	word_encountered = False  # Flag variable to track if the specific word has been encountered
	for line in content:
		modified_content.append(line)
		if specific_word in line and not word_encountered:
			modified_content.append(new_line_text + '\n')  # Add the new line of text
			word_encountered = True  # Set the flag to True after encountering the word

	# Write the modified content back to the tex file
	with open(tex_file, 'w', encoding='utf-8') as file:
		file.writelines(modified_content)

def extract_special_words(input_string):
	special_words = re.findall(r'\S*(?:fig:|_|\-)\S*', input_string)
	return special_words

def extract_paragraphs_by_style(docx_path, style_name):
	doc = Document(docx_path)
	paragraphs_with_style = []
	temp_files = []

	for i, paragraph in enumerate(doc.paragraphs):
		if paragraph.style.name == style_name:
			new_doc = Document()
			new_paragraph = new_doc.add_paragraph()

			# Copy run-level formatting from the original paragraph to the new paragraph
			for run in paragraph.runs:
				new_run = new_paragraph.add_run(run.text)
				new_run.bold = run.bold
				new_run.italic = run.italic
				new_run.underline = run.underline
				# ... copy other font properties as needed

			temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
			temp_files.append(temp_file.name)
			new_doc.save(temp_file.name)

			paragraphs_with_style.append(temp_file.name)

	return paragraphs_with_style

def replace_caption_cross_references(content):
	pattern = r'\\emph{\\textbf{\\ul{([^{}]+)}}}'

	def repl(match):
		cross_ref = match.group(1)
		cross_ref = cross_ref.replace("\\", "")
		return r'\ref{' + cross_ref + '}'

	modified_content = re.sub(pattern, repl, content)
	return modified_content




def handle_failed_chapter(e):
	print(bold_red, "ERROR: Failed to convert Chapter", " Error:", e, reset)


def replace_youtube_link_with_command(tex_file_path, website):
	try:
		# Read the content of the .tex file
		with open(tex_file_path, 'r', encoding='utf-8') as file:
			tex_content = file.read()
		
		if website == 'Youtube':
			# Define regular expressions for the YouTube links
			regex = r'https://youtu\.be/[^\s]+|https://www\.youtube\.com/[^\s]+'
		elif website == 'Vimeo':
			# Define regular expressions for the Vimeo links
			regex = r'https://vimeo\.com/[^\s]+'
		else:
			regex = r'https://youtu\.be/[^\s]+|https://www\.youtube\.com/[^\s]+'
			
		# Search for the first YouTube link in the file
		match = re.search(regex, tex_content)

		if match:
			# Extract the matched YouTube link
			website_link = match.group(0)
			short_link = re.sub(r'https://youtu\.be/', '', website_link)
			short_link = re.sub(r'https://youtube\.com/', '', short_link)
			short_link = re.sub(r'https://vimeo\.com/', '', short_link)

			if website == 'Youtube':
				# Format the YouTube link
				formatted_website_link = r'\youtube{' + short_link + '}'
			else:
				formatted_website_link = r'\vimeo{' + short_link + '}'
		else:
			return None, None
		
		# Find the entire unedited line where the URL was found
		lines = tex_content.split('\n')
		for i, line in enumerate(lines):
			if re.search(regex, line):
				unedited_line = lines[i]
				break
		else:
			unedited_line = None

		# Encode the modified text and unedited line into UTF-8 bytes with ASCII escape characters
		modified_bytes = formatted_website_link.encode('utf-8')
		unedited_bytes = unedited_line.encode('utf-8') if unedited_line is not None else None

		#return formatted_youtube_link, unedited_line
		return modified_bytes, unedited_bytes
	except Exception as e:
		return None, None


def replace_string_in_tex(tex_file, search_string, replace_string):
	# Define a custom replacement function
	def custom_replace(match):
		return replace_string

	# Read the content of the .tex file
	with open(tex_file, 'r', encoding='utf-8') as file:
		tex_content = file.read()

	# Perform the replacement using the custom function
	modified_content, replacements_count = re.subn(re.escape(search_string), custom_replace, tex_content, flags=re.MULTILINE)

	# Write the modified content back to the .tex file
	with open(tex_file, 'w', encoding='utf-8') as file:
		file.write(modified_content)

	return replacements_count


def get_docx_files(directory):
	docx_files = []
	
	for root, dirs, files in os.walk(directory):
		for file in files:
			if file.endswith(".docx"):
				file_path = os.path.join(root, file)
				docx_files.append(file_path)
	
	return docx_files

#***************** RUN FULL CONVERSION *********************************
#Pre convert locally for all files in current folder
docx_files = get_docx_files(current_directory)

for d in docx_files:
	print(d)
	try:
		gen_result = generate_tex(d)
		print("Finished Converting file:", d)
	except Exception as e:
		print("ERROR: Failed to convert", e)
print("Finished converting all files")



