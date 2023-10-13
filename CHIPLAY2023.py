#import os
from dataclasses import replace
import os
from shutil import ExecError
#from signal import pause
import zipfile
import tempfile
import pypandoc 
import pandocxnos
import re
import olsync
import docx
import docx2txt
import http.cookiejar
from docx import Document
from olsync import olclient, olbrowserlogin
from xml.etree import ElementTree 
from lxml import etree
from xml.etree.ElementTree import QName

# DEFINE GLOBALS *******************************
current_directory = os.getcwd()                                                                         
doc_path = current_directory 

      
def generate_tex(input_file):

    doc_file = input_file
    tex_file = input_file + ".tex"

      
    if not os.path.exists(doc_file):
        print("Missing file path: " + doc_file)
        # Handle the error condition appropriately
        return False
    else:
        print("Processing File:", doc_file)
        try:
                conversion_result = convert_docx_to_tex(doc_file, tex_file)
                if conversion_result:
                    try:
                            fixes_result = manual_fixes(input_file)
                            return fixes_result
                    except:
                            print("Manual Fixes failed")
        except :
                print("Atempted but Failed to convert document")
    return False


def convert_docx_to_tex(input_file, output_file):
      to_format = 'tex'
      extra_args = [
      '--filter', 'pandoc-fignos',
      #'--filter', 'pandoc-crossref',
      '--citeproc',
      #'--csl', 'custom-key.csl',
      '--natbib',
      '--bibliography', 'Bibliography.bib',
      #'--csl=author-year.csl', 
      #'--csl=acm-siggraph.csl',
      '--wrap', 'none',
      '--extract-media', './',       
      '--verbose'
      #'-H', 'path/to/caption.sty'  # Replace with the actual path to caption.sty
      ]
      input_format = 'docx+citations'
      try:
            pypandoc.convert_file(input_file, to_format, input_format, outputfile=output_file, extra_args=extra_args)
            return True
      except Exception as e:
            print("Pandoc ERROR:", str(e))
            return False

def manual_fixes(input_file):
      #create and verify zip version of document then unzip to temp
      doc_file = input_file
      tex_file = input_file + ".tex"
      result = False
      captions = []
      image_names = []

      try:
          image_names = get_original_image_names(doc_file)
    
      except Exception as e:
          print("Failed to find image names:", e)
          

      pattern = r'^\\includegraphics\[width=.*in,height=.*in\]{'
      try:
            captions = find_text_with_style(doc_file, "Caption")
      except Exception as e:
            print("Failed to find captions by style:", e)
      try:
            crossrefs = find_text_with_style(doc_file, "CrossReference")
      except Exception as e:
            print("Failed to find crossreferences by style:", e)

      try:      
            #remove empty captions
            for c in captions:
                  if c == "":
                        captions.remove(c)
      except Exception as e:
            print("Remove empty captions failed:", e)

      
      try:
            for name in image_names:
                  i = image_names.index(name)
                  has_note = False
                  search_pattern = ""
                  
                  try :
                        replace_line_with_pattern(tex_file, r"\\end{fignos:no-prefix-figure-caption}", "")
                        result = True
                  except:
                        print("Error changing to \\end{figure}")

                  try:
                        name = os.path.basename(name)
                  except Exception as e:
                        print(e)

                  try:      
                        replace_line_with_pattern(tex_file, r"\\begin{fignos:no-prefix-figure-caption}", r"\begin{figure}")
                        result = True
                  except:
                        print("Error changing \\begin{figure}")

                  label_name, *_ = name.split(".")
                  label_name = r"\label{fig:" + label_name + "}"
                  replacement = "\\includegraphics[width=\\textwidth]{" + "Figures/" + name + "}"
                  try:
                        replace_line_with_pattern(tex_file, pattern, replacement)
                        result = True    
                  except :
                        print("Error adjusting \\includegraphics")            

                  
                  try:#Figure titles
                        extra_args = [ 
                        '--biblatex',
                        '--citeproc',
                        '--wrap', 'none',
                        ]
                        latex_caption = pypandoc.convert_text(captions[i], 'latex', 'html', extra_args)
                        latex_caption = r"" + latex_caption
                        latex_caption = remove_trailing_whitespace(latex_caption)
                        search_pattern = get_first_4_words(latex_caption)
                        figure_title = ""
                        #Figure notes
                        delimiter = "Note."
                        if delimiter in latex_caption:
                            before_note, after_note = latex_caption.split(delimiter, 1)
                            after_note = delimiter + after_note # Include the delimiter in the "before_note" string
                            figure_title = r"\caption{" + before_note + "}"    #Figure Number and Long Title as stated in APA7
                            
                            caption_replacement = r"\small " + after_note
                            replace_line_with_pattern(tex_file, search_pattern, caption_replacement)
                            has_note = True
                        else:
                            print("String does not contain 'Note.'")
                            has_note = False
                            figure_title = r"\caption{" + latex_caption + "}"    #Figure Number and Long Title as stated in APA7
                            replace_line_with_pattern(tex_file, search_pattern, "")
                        try:
                            add_new_line_of_text_above_word(tex_file, replacement, figure_title)
                        except Exception as e:
                            print("Error with Figure title:", e)           
                        result = True                      
                  except Exception as e:
                        print("Failed Figure Title:", e)

                  try:#Label      
                        #cap_search = r"\\caption{" + search_pattern + ".*"
                        add_new_line_of_text_below_word(tex_file, replacement, label_name)
                        result = True
                  except Exception as e:
                        print("Error adding label:", e)            

                  try:#End Figure  
                        if has_note:
                            add_new_line_of_text_below_word(tex_file, caption_replacement, r"\end{figure}")
                        else:
                            add_new_line_of_text_below_word(tex_file, label_name, r"\end{figure}")
                        result = True
                  except Exception as e:
                        print("Error adding \\end{Figure} line:", e)
      except Exception as e:
            print("Image names loop failed:", e)
      try:
            replace_crossreferences(tex_file)
      except Exception as e:
            print("Replace captions failed:", e)
      for ref in crossrefs:
           try:
                 ref_search = r"\emph{\textbf{\ul{" + "ref" + "}}}"
                 ref_replacement = r"\ref{" + ref + "}"
                 find_and_replace(tex_file, ref_search, ref_replacement)
           except Exception as e:
                 print("Failed correcting crossref:", e)

      return result




def connect_to_overleaf():
      if use_cookie:
            # Create a CookieJar object
            cookie_jar = http.cookiejar.CookieJar()

            # Load the cookie from a file
            cookie_file = r"E:\Code\LaTexConversion\LaTexConversion\LaTexConversion\.olauth"
            try:
                  pass#http.cookiejar.CookieJar.
            except Exception as e:
                  print("Failed to Load cookie:", e)

            session =   {
                            "cookie": cookie_file,
                            "csrf": "token123"
                        }
      else:
            session = olbrowserlogin.login()
       
      if session is not None:
            client = olclient.OverleafClient(session["cookie"], session["csrf"])

            try:
                  #project= client.get_project(ol_project_name)
                  projects = client.all_projects()
                  print(projects)
            except Exception as e:
                  print("Error with acessing projects:", e)

            try:
                  project= client.get_project(ol_project_name)
            except Exception as e:
                  print("Error getting single project", e)
      else:
            print("Error with Overleaf Login")
      return overleaf_project



def sync_chapter_to_overleaf(project_ref, chapter_num):                  
      new_tex = update_file_number(chapter_num, ".tex")
      old_tex = "Chapter" + chapter_num + ".tex"
      project_ref.update_file(old_tex, new_tex)  


def update_file_number(chapter_number, extension):
       file = os.path.join(doc_path, 'Chapter')
       file = file + str(chapter_number) + extension
       return file

def get_original_image_names(docx_file):
      doc = Document(docx_file)
      image_names = []
      refs = []
      ordered_images = []
      all_rids = []

      try:
            for rel in doc.part.rels.values():
                  if "image" in rel.reltype:
                        try:
                              image_name = rel._target.split("/")[-1]
                              image_name = r"" + image_name
                              image_names.append(image_name)
                              ref = rel.rId 
                              print(ref)
                              refs.append(ref)
                        except Exception as e:
                              pass#print(e)
      except Exception as e:
            print("rel.values failed:", e)
      try:
            all_rids = get_rid_order(docx_file)
      except Exception as e:
            print("All rID's failed:",e)
      try:
            for rid in all_rids:
                  if rid in refs:
                        pass
                  else:
                        all_rids.remove(rid)
      except Exception as e:
            print("Fix all rID list fialed:",e)

      #reordered_rids = sorted(refs, key=lambda x: all_rids.index(x))
      try:
            ordered_images = [c for _, c in sorted(zip(refs, image_names), key=lambda x: all_rids.index(x[0]))]
      except Exception as e:
            print("order_images Failed:", e)
      return ordered_images

def remove_trailing_whitespace(string):
    return string.rstrip()

def get_first_4_words(string):
    words = string.split()
    first_4_words = words[:4]
    return ' '.join(first_4_words)

def add_line_above_first_line(tex_file, new_line):
    with open(tex_file, 'r') as file:
        lines = file.readlines()

    lines.insert(0, new_line + '\n')

    with open(tex_file, 'w') as file:
        file.writelines(lines)

def find_replace_unknown(original, replacement):
    # Escape special characters in the original string
    escaped_original = re.escape(original)
    
    # Create a regular expression pattern with the escaped original string
    pattern = re.compile(escaped_original)
    
    # Replace the unknown characters in the replacement string using the pattern
    modified = re.sub(pattern, replacement, original)
    
    return modified

def replace_crossreferences(file_path):
    pattern = r'\\emph{\\textbf{\\ul{([^{}]+)}}}'

    with open(file_path, 'r') as file:
        content = file.read()

    def repl(match):
          cross_ref = match.group(1)
          cross_ref = cross_ref.replace("\\", "")
          return r'\ref{' + cross_ref + '}'

    modified_content = re.sub(pattern, repl, content)

    with open(file_path, 'w') as file:
        file.write(modified_content)

def replace_first_occurrence(file_path, word_to_replace, replacement):
      with open(file_path, 'r') as file:
            content = file.read()

      # Find the index of the first occurrence of the word
      index = content.find(word_to_replace)

      if index != -1:
            # Replace the word with the desired replacement
            updated_content = content[:index] + replacement + content[index + len(word_to_replace):]

            # Write the updated content back to the file
            with open(file_path, 'w') as file:
                  file.write(updated_content)

def add_new_line_of_text_above_word(tex_file, specific_word, new_line_text):
    # Read the content of the tex file
    with open(tex_file, 'r') as file:
        content = file.readlines()

    # Iterate over the lines and add a new line of text above the first line with the specific word
    modified_content = []
    word_encountered = False  # Flag variable to track if the specific word has been encountered
    for line in content:
        if specific_word in line and not word_encountered:
            modified_content.append(new_line_text + '\n')  # Add the new line of text
            word_encountered = True  # Set the flag to True after encountering the word
        modified_content.append(line)

    # Write the modified content back to the tex file
    with open(tex_file, 'w') as file:
        file.writelines(modified_content)

def add_new_line_of_text_below_word(tex_file, specific_word, new_line_text):
    # Read the content of the tex file
    with open(tex_file, 'r') as file:
        content = file.readlines()

    # Iterate over the lines and add a new line of text above the first line with the specific word
    modified_content = []
    word_encountered = False  # Flag variable to track if the specific word has been encountered
    for line in content:
        modified_content.append(line)
        if specific_word in line and not word_encountered:
            modified_content.append(new_line_text + '\n')  # Add the new line of text
            word_encountered = True  # Set the flag to True after encountering the word

    # Write the modified content back to the tex file
    with open(tex_file, 'w') as file:
        file.writelines(modified_content)

def replace_line_with_pattern(file_path, pattern, replacement):
    replaced = False  # Flag to track if replacement has been made

    with open(file_path, 'r') as file:
        lines = file.readlines()

    with open(file_path, 'w') as file:
        for line in lines:
            if not replaced and re.match(pattern, line):
                file.write(replacement + '\n')
                replaced = True  # Set the flag to True after making the replacement
            else:
                file.write(line)

def add_line_above_pattern(file_path, pattern, new_line):
    with open(file_path, 'r') as file:
        content = file.read()

    modified_content = re.sub(pattern, new_line + r'\n\g<0>', content, count=1)

    with open(file_path, 'w') as file:
        file.write(modified_content)

def add_line_below_pattern(file_path, pattern, new_line):
    with open(file_path, 'r') as file:
        content = file.read()

    modified_content = re.sub(pattern, r'\g<0>\n' + new_line, content, count=1)

    with open(file_path, 'w') as file:
        file.write(modified_content)

def find_text_with_style(document_path, style_name):
    doc = Document(document_path)
    text_with_style = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_name:
            text_with_style.append(paragraph.text)

    return text_with_style

def find_and_replace(file_path, search_text, replace_text):
    with open(file_path, 'r') as file:
        content = file.read()

    try:
            modified_content = content.replace(search_text, replace_text)
            with open(file_path, 'w') as file:
                  file.write(modified_content)
    except Exception as e:
            print(e)

def get_rid_order(docx_file):
      # Open the document using python-docx
      doc = Document(docx_file)

      # Retrieve the document.xml part
      document_part = doc.part.document

      # Convert the document XML blob to string
      xml_string = document_part.part.blob.decode('utf-8')
      root = None
      drawing_elements = []
      namespaces = {}
      try:
            # Parse the document XML using lxml
            root = etree.fromstring(doc.part.blob, parser=None)
      except Exception as e:
            print("XML Parsing Failed:", e)

      try:
        # Register the 'w' namespace prefix
            namespaces = {
                'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
                'cx1': 'http://schemas.microsoft.com/office/drawing/2015/9/8/chartex',
                'cx2': 'http://schemas.microsoft.com/office/drawing/2015/10/21/chartex',
                'cx3': 'http://schemas.microsoft.com/office/drawing/2016/5/9/chartex',
                'cx4': 'http://schemas.microsoft.com/office/drawing/2016/5/10/chartex',
                'cx5': 'http://schemas.microsoft.com/office/drawing/2016/5/11/chartex',
                'cx6': 'http://schemas.microsoft.com/office/drawing/2016/5/12/chartex',
                'cx7': 'http://schemas.microsoft.com/office/drawing/2016/5/13/chartex',
                'cx8': 'http://schemas.microsoft.com/office/drawing/2016/5/14/chartex',
                'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
                'aink': 'http://schemas.microsoft.com/office/drawing/2016/ink',
                'am3d': 'http://schemas.micro...17/model3d',
                'o': 'urn:schemas-microsoft-com:office:office',
                'oel': 'http://schemas.microsoft.com/office/2019/extlst',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
                'v': 'urn:schemas-microsoft-com:vml',
                'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'w10': 'urn:schemas-microsoft-com:office:word',
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
                'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
                'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
                'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
                'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
                'w16du': 'http://schemas.microsoft.com/office/word/2023/wordml/word16du',
                'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
                'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
                'wpg': 'http://schemas.microsoft.com/office/' }

            
            # Find all w:drawing elements in the document.xml
            #drawing_elements = root.findall(".//w:drawing", namespaces=namespaces)
            if root != None:      
                  drawing_elements = root.findall(".//a:blip", namespaces=namespaces)
            else:
                 print("root is None")

      except Exception as e:
            print("w:drawing search failed:", e)

      # Retrieve the rId values in the order they appear in the document.xml
      rid_order = []
      try:
            for drawing_element in drawing_elements:
                  #blip_element = drawing_element.find(".//a:blip", namespaces=namespaces)
                  if drawing_element is not None:
                        embed_attrib_name = QName(namespaces['r'], 'link')
                        try:
                              str_attrib = str(embed_attrib_name)
                              r_id = drawing_element.attrib[str_attrib]
                              rid_order.append(r_id)
                        except KeyError:
                              print(KeyError)
      except Exception as e:
           print("rid list failed:", e)
      print("rids")
      print(rid_order)
      return rid_order

def get_docx_files(directory):
    docx_files = []
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                docx_files.append(file_path)
    
    return docx_files

#***************** RUN FULL CONVERSION *********************************
#Pre convert locally for all chapters
docx_files = get_docx_files(current_directory)

for d in docx_files:
    print(d)
    try:
        gen_result = generate_tex(d)
        print("Finished Converting file:", d)
    except Exception as e:
        print("ERROR: Failed to convert", e)
print("Finished converting all files")

