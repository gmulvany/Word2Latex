# By Gerard Mulvany
import os
import datetime
import tempfile
from xml.etree.ElementTree import QName
import re
import pypandoc
from docx import Document
from lxml import etree
import pandas as pd

####################### GLOBAL PROPERTIES ############################
REPLACE_VIDEO_LINKS = False
USE_APA_FIGURE_STYLE = True

current_directory = os.getcwd()

doc_path = current_directory
CUSTOM_TEMPLATE = "custom_template.latex"
# ANSI escape codes for console red text
BOLD_RED = "\033[1;31m"
GREEN = "\033[32m"
YELLOW = "\033[93m"
ITALIC = "\033[3m"
RESET = "\033[0m"

# Pandoc arguments
extra_args = [
    "--filter",
    "pandoc-fignos",
    "--citeproc",
    "--biblatex",
    "--wrap",
    "none",
    "--extract-media",
    "./",
    "--verbose",
]
TO_FORMAT = "tex"
INPUT_FORMAT = "docx+citations"

namespaces = {
    "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "cx": "http://schemas.microsoft.com/office/drawing/2014/chartex",
    "cx1": "http://schemas.microsoft.com/office/drawing/2015/9/8/chartex",
    "cx2": "http://schemas.microsoft.com/office/drawing/2015/10/21/chartex",
    "cx3": "http://schemas.microsoft.com/office/drawing/2016/5/9/chartex",
    "cx4": "http://schemas.microsoft.com/office/drawing/2016/5/10/chartex",
    "cx5": "http://schemas.microsoft.com/office/drawing/2016/5/11/chartex",
    "cx6": "http://schemas.microsoft.com/office/drawing/2016/5/12/chartex",
    "cx7": "http://schemas.microsoft.com/office/drawing/2016/5/13/chartex",
    "cx8": "http://schemas.microsoft.com/office/drawing/2016/5/14/chartex",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "aink": "http://schemas.microsoft.com/office/drawing/2016/ink",
    "am3d": "http://schemas.micro...17/model3d",
    "o": "urn:schemas-microsoft-com:office:office",
    "oel": "http://schemas.microsoft.com/office/2019/extlst",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w10": "urn:schemas-microsoft-com:office:word",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16du": "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wpg": "http://schemas.microsoft.com/office/",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

acm_preamble = [
    r"%%",
    r"%% This is file sample-manuscript.tex",
    r"%% generated with the docstrip utility.",
    r"%%",
    r"%% The original source files were:",
    r"%%",
    r"%% samples.dtx  (with options: manuscript)",
    r"%% ",
    r"%% IMPORTANT NOTICE:",
    r"%% ",
    r"%% For the copyright see the source file.",
    r"%% ",
    r"%% Any modified versions of this file must be renamed",
    r"%% with new filenames distinct from sample-manuscript.tex.",
    r"%% ",
    r"%% For distribution of the original source see the terms",
    r"%% for copying and modification in the file samples.dtx.",
    r"%% ",
    r"%% This generated file may be distributed as long as the",
    r"%% original source files, as listed above, are part of the",
    r"%% same distribution. (The sources need not necessarily be",
    r"%% in the same archive or directory.)",
    r"%%",
    r"%% Commands for TeXCount",
    r"%TC:macro \cite [option:text,text]",
    r"%TC:macro \citep [option:text,text]",
    r"%TC:macro \citet [option:text,text]",
    r"%TC:envir table 0 1",
    r"%TC:envir table* 0 1",
    r"%TC:envir tabular [ignore] word",
    r"%TC:envir displaymath 0 word",
    r"%TC:envir math 0 word",
    r"%TC:envir comment 0 0",
    r"%%",
    r"%%",
    r"%% The first command in your LaTeX source must be the \documentclass command.",
    r"\documentclass[manuscript,screen,review]{acmart}",
    r"",
    r"%%",
    r"%% \BibTeX command to typeset BibTeX logo in the docs",
    r"\AtBeginDocument{%",
    r"  \providecommand\BibTeX{{%",
    r"    \normalfont B\kern-0.5em{\scshape i\kern-0.25em b}\kern-0.8em\TeX}}}",
    r"",
    r"%% Rights management information.  This information is sent to you",
    r"%% when you complete the rights form.  These commands have SAMPLE",
    r"%% values in them; it is your responsibility as an author to replace",
    r"%% the commands and values with those provided to you when you",
    r"%% complete the rights form.",
    r"\setcopyright{acmcopyright}",
    r"\copyrightyear{2018}",
    r"\acmYear{2018}",
    r"\acmDOI{XXXXXXX.XXXXXXX}",
    r"",
    r"%% These commands are for a PROCEEDINGS abstract or paper.",
    r"\acmConference[Conference acronym XX]{Make sure to enter the correct",
    r"  conference title from your rights confirmation email}{June 03--05,",
    r"  2018}{Woodstock, NY}",
    r"\acmPrice{15.00}",
    r"\acmISBN{978-1-4503-XXXX-X/18/06}",
    r"",
    r"",
    r"%%",
    r"%% Submission ID.",
    r"%% Use this when submitting an article to a sponsored event. You'll",
    r"%% receive a unique submission ID from the organizers",
    r"%% of the event, and this ID should be used as the parameter to this command.",
    r"%%\acmSubmissionID{123-A56-BU3}",
    r"",
    r"%%",
    r"%% For managing citations, it is recommended to use bibliography",
    r"%% files in BibTeX format.",
    r"%%",
    r"%% You can then either use BibTeX with the ACM-Reference-Format style,",
    r"%% or BibLaTeX with the acmnumeric or acmauthoryear styles, that include",
    r"%% support for advanced citation of software artefact from the",
    r"%% biblatex-software package, also separately available on CTAN.",
    r"%%",
    r"%% Look at the sample-*-biblatex.tex files for templates showcasing",
    r"%% the biblatex styles.",
    r"%%",
    r"",
    r"%%",
    r"%% The majority of ACM publications use numbered citations and",
    r"%% references.  The command \citestyle{authoryear} switches to the",
    r'%% "author year" style.',
    r"%%",
    r"%% If you are preparing content for an event",
    r'%% sponsored by ACM SIGGRAPH, you must use the "author year" style of',
    r"%% citations and references.",
    r"%% Uncommenting",
    r"%% the next command will enable that style.",
    r"%%\citestyle{acmauthoryear}",
    r"",
    r"%%",
    r"%% end of the preamble, start of the body of the document source.",
    r"\begin{document}",
    r"",
    r"%%",
    r'%% The "title" command has an optional parameter,',
    r'%% allowing the author to define a "short title" to be used in page headers.',
]

title_commands = [
    r"%%",
    r"%% By default, the full list of authors will be used in the page",
    r"%% headers. Often, this list is too long, and will overlap",
    r"%% other information printed in the page headers. This command allows",
    r"%% the author to define a more concise list",
    r"%% of authors' names for this purpose.",
    r"",
    r"%%",
    r"%% The abstract is a short summary of the work to be presented in the",
    r"%% article.",
    r"\begin{abstract}",
    r"\end{abstract}",
    r"",
    r"%%",
    r"%% The code below is generated by the tool at http://dl.acm.org/ccs.cfm.",
    r"%% Please copy and paste the code instead of the example below.",
    r"%%",
    r"",
    r"%%",
    r"%% Keywords. The author(s) should pick words that accurately describe",
    r"%% the work being presented. Separate the keywords with commas.",
    r"\keywords{}",
    r"",
    r"\received{}",
    r"\received[revised]{}",
    r"\received[accepted]{}",
    r"",
    r"%%",
    r"%% This command processes the author and affiliation and title",
    r"%% information and builds the first part of the formatted document.",
    r"\maketitle",
]
####################################### MAIN FUNCTIONS ################################################################


def generate_tex(input_file):
    """Generate the LATEX file from a given docx file"""
    try:
        # update docx for current chapter number
        doc_file = input_file
        tex_file = input_file.replace(".docx", "")
        tex_file = tex_file + ".tex"
    except Exception as e:
        print(BOLD_RED, "Error: Find document failed :", e, RESET)
        return False

    if not os.path.exists(doc_file):
        print(BOLD_RED + "Error: Missing file path: " + doc_file + RESET)
        return False
    else:
        print("Processing File:", doc_file)
        try:
            conversion_result = convert_docx_to_tex(doc_file, tex_file)
            if conversion_result:
                try:
                    fixes_result = manual_fixes(input_file)
                    return fixes_result
                except Exception as e:
                    print(BOLD_RED, "Error: Manual Fixes failed:", e, RESET)
        except Exception as e:
            print(
                BOLD_RED, "Error: Attempted but failed to convert document:", e, RESET
            )
    return False


def convert_docx_to_tex(input_file, output_file):
    """Generate the LATEX file from a given docx file"""
    try:
        pypandoc.convert_file(
            input_file,
            TO_FORMAT,
            INPUT_FORMAT,
            outputfile=output_file,
            extra_args=extra_args,
        )
        return True
    except Exception as e:
        print(BOLD_RED, "Pandoc ERROR:", str(e), RESET)
        return False


def manual_fixes(input_file):
    """Manually fix many pandoc and latex issues with python"""
    # create and verify zip version of document then unzip to temp
    doc_file = input_file
    tex_file = input_file.replace(".docx", "")
    tex_file = tex_file + ".tex"
    result = False
    captions = []
    image_names = []
    subdoc_captions = []
    had_error = False
    try:
        image_names, image_alts = get_original_image_names(doc_file)
    except Exception as e:
        print(BOLD_RED, "Failed to find image names:", e, RESET)
        had_error = True
        # Copy each caption for figures into their own temporary docx so they can be individually pandoc converted and cross-references can be linked
    try:
        subdoc_captions = extract_paragraphs_by_style(doc_file, "Caption")
        #subdoc_tablecaps = extract_paragraphs_by_style(doc_file, "Table Caption")
    except Exception as e:
        print(BOLD_RED, "Failed to find captions by style:", e, RESET)
        had_error = True

    try:  # Get TITLE
        title_text = ""
        title_texts = find_text_with_style(doc_file, "Title")
        if title_texts.__len__() == 1:
            title_text = r"\title{" + title_texts[0] + "}"
    except Exception as e:
        print(YELLOW, "No Title or error with title", RESET)

    try:  # Get Abstract
        abstract_text = ""
        abstract_texts = []
        abstract_texts = extract_paragraphs_by_style(doc_file, "Abstract")
        if abstract_texts.__len__() == 1:
            abstract_text = pypandoc.convert_file(
                abstract_texts[0], TO_FORMAT, INPUT_FORMAT, extra_args=extra_args
            )
            abstract_text = abstract_text.encode("unicode-escape").decode("utf-8")
            abstract_text = abstract_text.replace("\\r", "").replace("\\n", "")
    except Exception as e:
        print(YELLOW, "No Title or error with title", RESET)

    # try:
    #     crossrefs = find_text_with_style(doc_file, "CrossReference")
    # except Exception as e:
    #     print(BOLD_RED, "Failed to find cross-references by style:", e, RESET)
    #     had_error = True

    try:  # remove empty captions
        for c in captions:
            if c == "":
                print(YELLOW, "Warning: Removed empty caption", RESET)
                had_error = True
                captions.remove(c)
    except Exception as e:
        print(BOLD_RED, "Remove empty captions failed:", e, RESET)
        had_error = True

    # Console line to inform
    print(GREEN, "Number of Images = ", len(image_names))
    print(GREEN, "Number of Captions = ", len(subdoc_captions), RESET)
    if len(image_names) != len(subdoc_captions):
        print(BOLD_RED, "ERROR: Figures and captions mismatch", RESET)
        had_error = True
    for name in image_names:
        i = image_names.index(name)
        description = image_alts[i]
        print(YELLOW, f"Progress: {i}/{len(image_names)}", RESET, end="\r")
        if USE_APA_FIGURE_STYLE:
            make_acm_figure(name, i, tex_file, subdoc_captions, description)
    try:
        replace_cross_references(
            tex_file
        )  # Run the function that replaces cross refs based on style with LATEX version
    except Exception as e:
        print(BOLD_RED, "Replace captions failed:", e, RESET)
        had_error = True

    for line in reversed(title_commands):
        add_line_above_first_line(tex_file, line)

    add_line_above_first_line(tex_file, title_text)
    for line in reversed(acm_preamble):
        add_line_above_first_line(tex_file, line)

    add_line_below_last_line(tex_file, "")
    add_line_below_last_line(tex_file, "")
    add_line_below_last_line(tex_file, r"\end{document}")
    add_line_below_last_line(tex_file, r"\endinput")

    current_year = datetime.datetime.now().year
    replace_line_with_pattern(
        tex_file, "\\acmYear{", r"\acmYear{" + current_year.__str__() + "}"
    )
    replace_line_with_pattern(
        tex_file, "copyrightyear{", "\\copyrightyear{" + current_year.__str__() + "}"
    )

    all_author_info = get_author_info()
    add_line_below_pattern(tex_file, r"of authors' names for this purpose.", "}")
    for author in reversed(all_author_info):
        if author["postcode"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \postcode{" + author["postcode"].__str__() + "}",
            )
        if author["country"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \country{" + author["country"].__str__() + "}",
            )
        if author["state"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \state{" + author["state"].__str__() + "}",
            )
        if author["city"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \city{" + author["city"].__str__() + "}",
            )
        if author["street"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \streetaddress{" + author["street"].__str__() + "}",
            )
        if author["institution"] != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"    \institution{" + author["institution"].__str__() + "}",
            )
        add_line_below_pattern(
            tex_file, r"of authors' names for this purpose.", r"\affiliation{%"
        )
        if author["email"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"\email{" + author["email"].__str__() + "}",
            )
        if author["mark"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"\authornotemark{" + author["mark"].__str__() + "}",
            )
        if author["orcid"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"\orcid{" + author["orcid"].__str__() + "}",
            )
        if author["note"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"\authornote{" + author["note"].__str__() + "}",
            )
        if author["name"].__str__() != "nan":
            add_line_below_pattern(
                tex_file,
                r"of authors' names for this purpose.",
                r"\author{" + author["name"].__str__() + "}",
            )

    if abstract_texts.__len__() == 1:
        # add_line_below_pattern(tex_file, r"\begin{abstract}", abstract_text)
        add_new_line_of_text_below_word(tex_file, r"\begin{abstract}", abstract_text)

    if len(image_names) == 0 and had_error == False:
        return True
    result = not had_error
    return result


def get_author_info():
    # Load the Excel file
    excel_file = "Authors.xlsx"
    df = pd.read_excel(excel_file)
    all_authors = []

    # Iterate through rows
    for index, row in df.iterrows():
        # Check if the row has content
        if not row.empty:
            author_data = {
                "name": "",
                "note": "",
                "email": "",
                "orcid": "",
                "mark": "",
                "institution": "",
                "street": "",
                "city": "",
                "state": "",
                "country": "",
                "postcode": "",
            }
            # Access and save column values as variables
            author_data["name"] = row["NAME"]
            author_data["note"] = row["NOTES"]
            author_data["email"] = row["EMAIL"]
            author_data["orcid"] = row["ORCID"]
            author_data["mark"] = row["NOTE_MARK"]
            author_data["institution"] = row["Institution"]
            author_data["street"] = row["Street Address"]
            author_data["city"] = row["City"]
            author_data["state"] = row["State"]
            author_data["country"] = row["Country"]
            author_data["postcode"] = row["Postcode"]
            all_authors.append(author_data)
    return all_authors


def make_acm_figure(name, i, tex_file, subdoc_captions, description):
    """Modifies figures to match ACM submission standards"""
    saved_note = ""
    try:
        replace_line_with_pattern(
            tex_file, r"\\end{fignos:no-prefix-figure-caption}", ""
        )
        replace_line_with_pattern(
            tex_file, r"\\begin{fignos:no-prefix-figure-caption}", r"\begin{figure}"
        )
    except Exception as e:
        print(BOLD_RED, e, RESET)
        return False
    description = r"    \Description{" + description + "}"
    name = os.path.basename(name)
    label_name, *_ = name.split(".")
    label_name = r"    \label{fig:" + label_name + "}"
    replacement = "    \\includegraphics[width=\\textwidth]{" + "Figures/" + name + "}"
    try:  # Make figures text width and add image filepath
        replace_line_with_pattern(
            tex_file, r"^\\includegraphics\[width=.*in,height=.*in\]{", replacement
        )
    except Exception as e:
        print(BOLD_RED, "Error adjusting \\includegraphics: ", e, RESET)
        return False
    latex_caption = pypandoc.convert_file(
        subdoc_captions[i], TO_FORMAT, INPUT_FORMAT, extra_args=extra_args
    )
    latex_caption = replace_caption_cross_references(latex_caption)
    latex_caption = latex_caption.encode("unicode-escape").decode("utf-8")
    latex_caption = latex_caption.replace("\\r", "").replace("\\n", "")
    saved_note = latex_caption
    latex_caption = r"    \caption{" + latex_caption + "}"
    # latex_caption = latex_caption.decode("utf-8")
    search_pattern = truncate_and_encode(saved_note, 50)
    saved_note = replace_line_with_pattern(
        tex_file, search_pattern, ""
    )  # Replace the original under-image caption with only the "after-note" section
    add_new_line_of_text_below_word(tex_file, replacement, latex_caption)
    add_new_line_of_text_above_word(tex_file, replacement, r"    \centering")
    add_new_line_of_text_below_word(tex_file, latex_caption, label_name)
    add_new_line_of_text_below_word(
        tex_file, label_name, description
    )  # Add ALT text description for accessability
    add_new_line_of_text_below_word(
        tex_file, description, r"\end{figure}"
    )  # Close the figure environment
    return True


def create_video_icons(tex_file):
    """Replaces youtube and vimeo links with small icons instead"""
    try:
        modified_text, unedited_line = replace_youtube_link_with_command(
            tex_file, "Youtube"
        )
        if modified_text is not None:
            add_new_line_of_text_above_word(
                tex_file,
                unedited_line.decode("utf-8", "replace"),
                modified_text.decode("utf-8", "replace")[:-1],
            )
    except Exception as e:
        print(BOLD_RED, e, RESET)
    try:
        modified_text, unedited_line = replace_youtube_link_with_command(
            tex_file, "Vimeo"
        )
        if modified_text is not None:
            add_new_line_of_text_above_word(
                tex_file,
                unedited_line.decode("utf-8", "replace"),
                modified_text.decode("utf-8", "replace")[:-1],
            )
    except Exception as e:
        print(BOLD_RED, e, RESET)


######################### SHORT FUNCTIONS ############################

def get_original_image_names(docx_file):
    """Returns the file names and alt-text descriptions for all images present 
    in the word document and orders them based on their order of appearance"""
    doc = Document(docx_file)
    image_names = []
    image_alts = []
    refs = []
    ordered_images = []
    all_rids = []
    try:
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                target_part = doc.part.rels[rel].target_part

                if target_part.content_type.startswith("image/"):
                    image_name = rel
                    next_rel = increment_rel(rel)
                for val in doc.part.rels.values():
                    if "image" in val.reltype and val.rId == next_rel:
                        image_name = val._target.split("/")[-1]
                image_name = r"" + image_name
                image_names.append(image_name)
                ref = next_rel
                refs.append(ref)
                main_document_xml = doc.part.blob
                root = etree.fromstring(main_document_xml)
                search = ".//a:blip[@r:embed='{}']".format(rel)
                pic_elem = root.find(search, namespaces=namespaces)
                if pic_elem is not None:
                    description = root.find(
						".//pic:pic/pic:nvPicPr/pic:cNvPr", namespaces=namespaces
					).get("descr")
                    image_alts.append(description)
                search = ".//a:blip[@r:embed='{}']".format(next_rel)
                pic_elem = root.find(search, namespaces=namespaces)
                if pic_elem is not None:
                    description = root.find(
						".//pic:pic/pic:nvPicPr/pic:cNvPr", namespaces=namespaces
					).get("descr")
                    image_alts.append(description)

    except Exception as e:
        print(BOLD_RED, "rel.values failed:", e, RESET)
    try:
        all_rids = get_rid_order(docx_file)
    except Exception as e:
        print(BOLD_RED, "All rID's failed:", e, RESET)
    try:
        for rid in all_rids:
            if rid in refs:
                pass
            else:
                all_rids.remove(rid)
    except Exception as e:
        print(BOLD_RED, "Fix all rID list failed:", e, RESET)

    # reordered_rids = sorted(refs, key=lambda x: all_rids.index(x))
    try:
        ordered_images = [
            c
            for _, c in sorted(
                zip(refs, image_names), key=lambda x: all_rids.index(x[0])
            )
        ]
        ordered_descriptions = [
            c
            for _, c in sorted(
                zip(refs, image_alts), key=lambda x: all_rids.index(x[0])
            )
        ]
    except Exception as e:
        print(BOLD_RED, "order_images Failed:", e, RESET)
    return ordered_images, ordered_descriptions


def increment_rel(rId_str, increment=1):
    """Increments the rId string by the given number (Default is by 1)"""
    # Use regular expressions to find the number part of the string
    match = re.search(r"\d+$", rId_str)

    if match:
        # Extract the number from the matched group
        number = int(match.group())
        # Increment the number
        new_number = number + increment
        # Replace the number in the original string with the incremented number
        return re.sub(r"\d+$", str(new_number), rId_str)

    # If no number was found, return the original string
    return rId_str


def truncate_and_encode(input_string, max_length=30):
    """Returns the given number of characters and encodes in binary form utf-8"""
    # Check if the input string length is less than or equal to max_length
    if len(input_string) <= max_length:
        # Encode the entire string to UTF-8 with escape characters
        if isinstance(input_string, bytes):
            return input_string
        return input_string.encode("utf-8", "escape")
    else:
        if isinstance(input_string, bytes):
            input_string = input_string.decode("utf-8")
        # Truncate the string to the first 30 characters and encode to UTF-8 with escape characters
        return input_string[:max_length].encode("utf-8", "escape")


def add_line_above_first_line(tex_file, new_line):
    """Add a new line before the current first"""
    with open(tex_file, "r", encoding="utf-8") as file:
        lines = file.readlines()

    lines.insert(0, new_line + "\n")

    with open(tex_file, "w", encoding="utf-8") as file:
        file.writelines(lines)


def add_line_below_last_line(tex_file, new_line):
    """Add a new line at the very bottom of the document"""
    with open(tex_file, "r", encoding="utf-8") as file:
        lines = file.readlines()

    lines.append(new_line + "\n")

    with open(tex_file, "w", encoding="utf-8") as file:
        file.writelines(lines)

def replace_cross_references(file_path):
    """Converts the given style in Word into LATEX cross references"""
    pattern = r"\\emph{\\textbf{\\ul{([^{}]+)}}}"

    with open(file_path, "r", encoding="utf-8") as file:
        content = file.read()

    def repl(match):
        cross_ref = match.group(1)
        cross_ref = cross_ref.replace("\\", "")
        return r"\ref{" + cross_ref + "}"

    modified_content = re.sub(pattern, repl, content)

    with open(file_path, "w", encoding="utf-8") as file:
        file.write(modified_content)

def replace_line_with_pattern(file_path, pattern, replacement):
    """Replaces a given line of tex in a LATEX document with the input replacement"""
    replaced = False  # Flag to track if replacement has been made
    after_note = ""
    if isinstance(replacement, bytes):
        replacement = replacement.decode("utf-8")

    # Check if 'pattern' is of type bytes
    if isinstance(pattern, bytes):
        pattern = pattern.decode("utf-8")

    with open(file_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    with open(file_path, "w", encoding="utf-8") as file:
        for line in lines:
            # Check if 'pattern' is of type bytes
            if isinstance(line, bytes):
                line = line.decode("utf-8")
            if not replaced:
                if pattern in line or re.search(pattern, line):
                    file.write(replacement + "\n")
                    replaced = True  # Set the flag to True after making the replacement
                else:
                    file.write(line)
            else:
                file.write(line)
    return after_note


def add_line_below_pattern(file_path, pattern, new_line):
    """Add a new line of text below the given pattern"""
    with open(file_path, "rb") as file:
        content = file.read()

    # Ensure that the pattern is a bytes object
    if isinstance(pattern, str):
        pattern = pattern.encode("utf-8")

    # Use bytes pattern to search in content
    pattern = re.compile(pattern)
    modified_content = pattern.sub(
        lambda x: x.group(0) + b"\n" + new_line.encode("utf-8"), content, count=1
    )

    with open(file_path, "wb") as file:
        file.write(modified_content)


def find_text_with_style(document_path, style_name):
    """Returns a string list for a given Word style"""
    doc = Document(document_path)
    text_with_style = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_name:
            text_with_style.append(paragraph.text)

    return text_with_style


def get_rid_order(docx_file):
    """Returns the order in which each rId appears in the Word document"""
    # Open the document using python-docx
    doc = Document(docx_file)

    # Retrieve the document.xml part
    document_part = doc.part.document

    # Convert the document XML blob to string
    # xml_string = document_part.part.blob.decode('utf-8')
    root = None
    drawing_elements = []
    try:
        # Parse the document XML using lxml
        root = etree.fromstring(doc.part.blob, parser=None)
    except Exception as e:
        print(BOLD_RED, "XML Parsing Failed:", e, RESET)

    try:
        # Register the 'w' namespace prefix

        # Find all w:drawing elements in the document.xml
        # drawing_elements = root.findall(".//w:drawing", namespaces=namespaces)
        if root != None:
            drawing_elements = root.findall(".//a:blip", namespaces=namespaces)
        else:
            print(BOLD_RED, "Get RID Order Failed: doc.part.blob.root is None", RESET)

    except Exception as e:
        print(BOLD_RED, "w:drawing search failed:", e, RESET)

    # Retrieve the rId values in the order they appear in the document.xml
    rid_order = []
    try:
        for drawing_element in drawing_elements:
            # blip_element = drawing_element.find(".//a:blip", namespaces=namespaces)
            if drawing_element is not None:
                embed_attrib_name = QName(namespaces["r"], "link")
                try:
                    str_attrib = str(embed_attrib_name)
                    r_id = drawing_element.attrib[str_attrib]
                    rid_order.append(r_id)
                except KeyError:
                    print(BOLD_RED, "Failed to get image link:", KeyError, RESET)
    except Exception as e:
        print(BOLD_RED, "rid list failed with a non-key-error:", e)
    return rid_order


def add_new_line_of_text_above_word(tex_file, specific_word, new_line_text):
    """Adds a new line of text above a line containing the given pattern"""
    # Read the content of the tex file
    with open(tex_file, "r", encoding="utf-8") as file:
        content = file.readlines()

    # Iterate over the lines and add a new line of text above the first line with the specific word
    modified_content = []
    word_encountered = (
        False  # Flag variable to track if the specific word has been encountered
    )
    for line in content:
        if specific_word in line and not word_encountered:
            modified_content.append(new_line_text + "\n")  # Add the new line of text
            word_encountered = True  # Set the flag to True after encountering the word
        modified_content.append(line)

    # Write the modified content back to the tex file
    with open(tex_file, "w", encoding="utf-8") as file:
        file.writelines(modified_content)


def add_new_line_of_text_below_word(tex_file, specific_word, new_line_text):
    """Adds a new line of text below a line containing the given pattern"""
    # Read the content of the tex file
    with open(tex_file, "r", encoding="utf-8") as file:
        content = file.readlines()

    # Iterate over the lines and add a new line of text above the first line with the specific word
    modified_content = []
    word_encountered = (
        False  # Flag variable to track if the specific word has been encountered
    )
    for line in content:
        modified_content.append(line)
        if isinstance(specific_word, bytes):
            specific_word = specific_word.decode("utf-8")
        if specific_word in line and not word_encountered:
            if isinstance(new_line_text, bytes):
                new_line_text = new_line_text.decode("utf-8")
            modified_content.append(new_line_text + "\n")  # Add the new line of text
            word_encountered = True  # Set the flag to True after encountering the word

    # Write the modified content back to the tex file
    with open(tex_file, "w", encoding="utf-8") as file:
        file.writelines(modified_content)


def extract_paragraphs_by_style(docx_path, style_name):
    """Creates temporary Word document containing paragraphs of a given style"""
    doc = Document(docx_path)
    paragraphs_with_style = []
    temp_files = []

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == style_name:
            new_doc = Document()
            new_paragraph = new_doc.add_paragraph()

            # Copy run-level formatting from the original paragraph to the new paragraph
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                # ... copy other font properties as needed

            temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_files.append(temp_file.name)
            new_doc.save(temp_file.name)

            paragraphs_with_style.append(temp_file.name)

    return paragraphs_with_style


def replace_caption_cross_references(content):
    """Converts the given style into LATEX cross-references that are within captions"""
    pattern = r"\\emph{\\textbf{\\ul{([^{}]+)}}}"

    def repl(match):
        cross_ref = match.group(1)
        cross_ref = cross_ref.replace("\\", "")
        return r"\ref{" + cross_ref + "}"

    modified_content = re.sub(pattern, repl, content)
    return modified_content



def replace_youtube_link_with_command(tex_file_path, website):
    """Replaces youtube and vimeo links with icons. Only works with my LATEX class"""
    try:
        # Read the content of the .tex file
        with open(tex_file_path, "r", encoding="utf-8") as file:
            tex_content = file.read()

        if website == "Youtube":
            # Define regular expressions for the YouTube links
            regex = r"https://youtu\.be/[^\s]+|https://www\.youtube\.com/[^\s]+"
        elif website == "Vimeo":
            # Define regular expressions for the Vimeo links
            regex = r"https://vimeo\.com/[^\s]+"
        else:
            regex = r"https://youtu\.be/[^\s]+|https://www\.youtube\.com/[^\s]+"

        # Search for the first YouTube link in the file
        match = re.search(regex, tex_content)

        if match:
            # Extract the matched YouTube link
            website_link = match.group(0)
            short_link = re.sub(r"https://youtu\.be/", "", website_link)
            short_link = re.sub(r"https://youtube\.com/", "", short_link)
            short_link = re.sub(r"https://vimeo\.com/", "", short_link)

            if website == "Youtube":
                # Format the YouTube link
                formatted_website_link = r"\youtube{" + short_link + "}"
            else:
                formatted_website_link = r"\vimeo{" + short_link + "}"
        else:
            return None, None

        # Find the entire unedited line where the URL was found
        lines = tex_content.split("\n")
        for i, line in enumerate(lines):
            if re.search(regex, line):
                unedited_line = lines[i]
                break
        else:
            unedited_line = None

        # Encode the modified text and unedited line into UTF-8 bytes with ASCII escape characters
        modified_bytes = formatted_website_link.encode("utf-8")
        unedited_bytes = (
            unedited_line.encode("utf-8") if unedited_line is not None else None
        )

        # return formatted_youtube_link, unedited_line
        return modified_bytes, unedited_bytes
    except Exception as e:
        return None, None

def get_docx_files(directory):
    """Returns all the .docx files in the current directory"""
    word_docx_files = []

    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                word_docx_files.append(file_path)

    return word_docx_files


# ***************** RUN FULL CONVERSION *********************************
# Pre convert locally for all files in current folder
docx_files = get_docx_files(current_directory)

for d in docx_files:
    print(d)
    try:
        GEN_RESULT = generate_tex(d)
        print("Finished Converting file:", d)
    except Exception as e:
        print("ERROR: Failed to convert", e)
print("Finished converting all files")
